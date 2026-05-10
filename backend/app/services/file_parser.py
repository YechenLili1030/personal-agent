"""文件解析服务 — PDF / Word(含文本框) / Excel / TXT / Markdown / 图片(多模态)"""

import base64
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls",
    ".txt", ".md", ".markdown",
    ".png", ".jpg", ".jpeg", ".bmp", ".gif",
}

SCANNED_PDF_THRESHOLD = 100  # 少于100字符视为扫描件


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "pdf", ".docx": "docx",
        ".xlsx": "xlsx", ".xls": "xlsx",
        ".txt": "txt", ".md": "md", ".markdown": "md",
        ".png": "image", ".jpg": "image", ".jpeg": "image",
        ".bmp": "image", ".gif": "image",
    }
    return mapping.get(ext, "unknown")


class ParseResult:
    def __init__(self, text: str, needs_multimodal: bool = False, structure: str = "semantic"):
        self.text = text
        self.needs_multimodal = needs_multimodal  # 需要多模态模型辅助
        self.structure = structure  # semantic / excel / markdown


def parse_file(file_path: str, file_type: str) -> ParseResult:
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "xlsx": _parse_xlsx,
        "txt": _parse_txt,
        "md": _parse_md,
        "image": _parse_image,
    }
    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")
    return parser(file_path)


# =========================== PDF ===========================

def _parse_pdf(path: str) -> ParseResult:
    texts = []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        result = "\n\n".join(texts)
        if len(result.strip()) >= SCANNED_PDF_THRESHOLD:
            return ParseResult(result, structure="semantic")
    except Exception:
        pass

    from pypdf import PdfReader
    reader = PdfReader(path)
    for page in reader.pages:
        t = page.extract_text()
        if t:
            texts.append(t)
    result = "\n\n".join(texts)

    if len(result.strip()) < SCANNED_PDF_THRESHOLD:
        logger.info("PDF 文本量过少(%d字)，标记为扫描件: %s", len(result.strip()), path)
        return ParseResult("", needs_multimodal=True, structure="semantic")

    return ParseResult(result, structure="semantic")


# =========================== Word ===========================

def _parse_docx(path: str) -> ParseResult:
    from docx import Document

    doc = Document(path)
    lines = []

    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())

    # 文本框
    body = doc.element.body
    for txbx in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
        text = _extract_text_from_element(txbx)
        if text.strip():
            lines.append(text.strip())
    for txbx in body.iter('{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'):
        text = _extract_text_from_element(txbx)
        if text.strip():
            lines.append(text.strip())

    return ParseResult("\n".join(lines), structure="semantic")


def _extract_text_from_element(el) -> str:
    texts = []
    for t in el.iter():
        if t.tag.endswith('}t') and t.text:
            texts.append(t.text)
        if t.tag.endswith('}tab') and t.tail:
            texts.append(t.tail)
    return ''.join(texts)


# =========================== Excel ===========================

def _parse_xlsx(path: str) -> ParseResult:
    """保留结构化格式，供后续 Excel 专用分块"""
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"## Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) if v is not None else "" for v in row]
            if any(vals):
                lines.append("\t".join(vals))
    return ParseResult("\n".join(lines), structure="excel")


# =========================== Markdown ===========================

def _parse_md(path: str) -> ParseResult:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return ParseResult(f.read(), structure="markdown")


# =========================== TXT ===========================

def _parse_txt(path: str) -> ParseResult:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return ParseResult(f.read(), structure="semantic")


# =========================== 图片 ===========================

def _parse_image(path: str) -> ParseResult:
    return ParseResult("", needs_multimodal=True, structure="semantic")


# =========================== 多模态辅助 ===========================

def encode_image_base64(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def get_image_mime(path: str) -> str:
    ext = Path(path).suffix.lower()
    mimes = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
             ".gif": "image/gif", ".bmp": "image/bmp"}
    return mimes.get(ext, "image/png")


def encode_pdf_pages_as_images(path: str) -> list[str]:
    """将 PDF 每页渲染为图片并返回 base64 列表"""
    try:
        import pdfplumber
        from PIL import Image
        images = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                img = page.to_image(resolution=150)
                buf = io.BytesIO()
                img.original.save(buf, format="PNG")
                images.append(base64.b64encode(buf.getvalue()).decode("utf-8"))
        return images
    except Exception as e:
        logger.error("PDF 渲染图片失败: %s", e)
        return []
